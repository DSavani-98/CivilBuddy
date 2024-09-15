from flask import Flask, request, send_file, jsonify, session
from pymongo import MongoClient
import os
from io import BytesIO
from PIL import Image
from bson import ObjectId
import gridfs
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gridfs.errors import NoFile
import re
import base64
import dotenv
import requests

dotenv.load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

app = Flask(__name__, static_folder='../frontend/build', static_url_path='/')
app.secret_key = 'your_secret_key'  # Required for session management

mongo_host = 'mongodb'
mongo_uri = f'mongodb://{mongo_host}:27017/'
client = MongoClient(mongo_uri)
db = client['flask_db']
fs = gridfs.GridFS(db)

# Serve the React App
@app.route('/')
def index():
    return app.send_static_file('index.html')

# Handle file upload
@app.route('/upload', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return jsonify({'error': 'No image uploaded'}), 400

    image = request.files['image']
    img_name = image.filename
    img_name_without_ext = re.sub(r'\.[^.]+$', '', img_name)  # Remove the extension
    
    # Get the Base64 encoding of the image
    image.seek(0)  # Make sure the file pointer is at the start
    base64_image = base64.b64encode(image.read()).decode('utf-8')

    # Generate report on the input image using OpenAI's gpt-4o-mini model
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    } 
    payload = {
        "model": "gpt-4o-mini",
        "messages": [
            {
                "role": "system",
                "content": "You possess advanced qualifications in civil engineering and architecture, with extensive expertise in the construction of various types of structures and buildings. With over 20 years of experience in producing technical reports, you excel in analyzing structural drawings and images. Your reports are crafted in a formal, technical style, ensuring precise and professional content."
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Write a technical report on the following image/drawing."
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 1000,
        "temperature": 0.3
    }

    response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
    result = response.json()
    report = result["choices"][0]["message"]["content"]

    # Save image in MongoDB
    img_id = fs.put(image, filename=img_name)

    # Create a Word document and format it
    document = Document()

    # 1. Add Title
    title = document.add_heading('Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Add Image to the Document
    img = Image.open(image)
    max_width_in_inches = 6  # You can adjust the maximum width if needed
    aspect_ratio = img.size[1] / img.size[0]
    new_width = min(Inches(max_width_in_inches), Inches(img.size[0]))
    new_height = Inches(aspect_ratio * new_width / Inches(1))

    # Save the image temporarily to add to the Word document
    temp_image_path = f'/tmp/{img_name}'
    img.save(temp_image_path)

    # Add image to Word document
    document.add_picture(temp_image_path, width=new_width, height=new_height)

    # 3. Add Image Caption (Filename)
    caption_paragraph = document.add_paragraph(img_name)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the caption

    # 4. Add the Generated Report Content
    document.add_paragraph(report)

    # Save the Word document as "Report_<filename>.docx"
    docx_filename = f'Report_{img_name_without_ext}.docx'
    temp_docx_path = f'/tmp/{docx_filename}'
    document.save(temp_docx_path)

    # Store the Word document in MongoDB
    with open(temp_docx_path, 'rb') as docx_file:
        docx_id = fs.put(docx_file, filename=docx_filename)

    # Remove temporary files
    os.remove(temp_image_path)
    os.remove(temp_docx_path)

    # Store docx_id in the session for user-specific access
    session['docx_id'] = str(docx_id)

    return jsonify({'image_id': str(img_id), 'docx_id': str(docx_id)}), 201

# Handle file download, ensuring only the file uploaded by the user can be downloaded
@app.route('/download/<docx_id>', methods=['GET'])
def download_docx(docx_id):
    if 'docx_id' not in session or session['docx_id'] != docx_id:
        return jsonify({'error': 'You are not authorized to download this file'}), 403

    try:
        # Fetch the document from MongoDB
        doc = fs.get(ObjectId(docx_id))
        return send_file(BytesIO(doc.read()), download_name=doc.filename, as_attachment=True)
    except NoFile:
        return jsonify({'error': f'File not found in GridFS with the given ID: {docx_id}'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)